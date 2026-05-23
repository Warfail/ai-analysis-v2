# app.py - FIXED WITH DEPARTMENT EXTRACTION
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, PlainTextResponse, HTMLResponse
import sqlite3
from datetime import datetime, timedelta
import io
import json
import re
from collections import Counter
import os
import traceback
import csv
import PyPDF2
from docx import Document
from database import execute_query, get_connection
# Initialize database on startup
from database import init_db
init_db()
# In app.py, at the top with other imports:
from database import get_connection, execute_query, fetch_all, fetch_one, insert
app = FastAPI(title="Report Analyzer AI - Fixed", version="2.2")

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==================== ENHANCED FILE PROCESSOR ====================
class EnhancedFileProcessor:
    """Process all file types: PDF, Word, Excel, CSV, Text"""
    
    def __init__(self):
        print("✅ EnhancedFileProcessor initialized")
    
    def process_file(self, content_bytes: bytes, filename: str) -> dict:
        filename_lower = filename.lower()
        
        print(f"📁 Processing {filename} ({len(content_bytes)} bytes)")
        
        result = {
            "success": True,
            "filename": filename,
            "file_size": len(content_bytes),
            "file_type": "unknown",
            "content": "",
            "structured_data": {},
            "error": None
        }
        
        try:
            if filename_lower.endswith('.csv'):
                return self._process_csv(content_bytes, result)
            elif filename_lower.endswith('.pdf'):
                return self._process_pdf(content_bytes, result)
            elif filename_lower.endswith(('.docx', '.doc')):
                return self._process_word(content_bytes, result)
            elif filename_lower.endswith(('.xlsx', '.xls')):
                return self._process_excel(content_bytes, result)
            elif filename_lower.endswith('.txt'):
                return self._process_text(content_bytes, result)
            else:
                return self._process_generic(content_bytes, result, filename_lower)
                
        except Exception as e:
            result["success"] = False
            result["error"] = str(e)
            result["content"] = f"Error processing file: {str(e)}"
            return result
    
    def _process_text(self, content_bytes: bytes, result: dict) -> dict:
        """Process text files"""
        try:
            content = content_bytes.decode('utf-8', errors='ignore')
            result["file_type"] = "text"
            result["content"] = content[:50000]
            
            # Try to extract structured data from text
            structured = self._extract_structured_from_text(content)
            if structured.get("has_structured_data"):
                result["structured_data"] = structured
            
            return result
        except Exception as e:
            result["success"] = False
            result["error"] = str(e)
            return result
    
    def _process_csv(self, content_bytes: bytes, result: dict) -> dict:
        """Process CSV files - FIXED TO ACCEPT EMPTY TEMPLATE"""
        try:
            content = content_bytes.decode('utf-8-sig')  # Handle UTF-8 BOM
            
            # DEBUG: Log first few lines
            lines = content.split('\n')
            print(f"CSV lines: {len(lines)}")
            print(f"First line: {lines[0] if lines else 'Empty'}")
            
            # Check if it's an empty template (just headers)
            if len(lines) <= 1:
                print("⚠️ CSV has only headers or is empty - treating as template")
                result["file_type"] = "csv_template"
                result["content"] = "Empty CSV template (headers only)"
                result["structured_data"] = {
                    "has_structured_data": False,
                    "message": "This appears to be a template. Add data to rows for analysis."
                }
                return result
            
            # Try 5-column structured parsing
            structured_5col = self._parse_5column_csv(content)
            if structured_5col and structured_5col.get("has_structured_data"):
                result["file_type"] = "structured_csv"
                result["content"] = structured_5col.get("formatted_content", content[:10000])
                result["structured_data"] = structured_5col
                return result
            
            # Try generic CSV parsing
            structured_generic = self._parse_generic_csv(content)
            if structured_generic and structured_generic.get("has_structured_data"):
                result["file_type"] = "csv"
                result["content"] = structured_generic.get("formatted_content", content[:10000])
                result["structured_data"] = structured_generic
                return result
            
            # Fall back to simple CSV display
            result["file_type"] = "csv"
            result["content"] = content[:10000]
            return result
        except Exception as e:
            print(f"CSV processing error: {e}")
            result["success"] = False
            result["error"] = f"CSV processing error: {str(e)}"
            return result
    
    def _parse_5column_csv(self, content: str) -> dict:
        """Parse 5-column CSV: Department, Date, Problems, Achievements, Future_Plan"""
        try:
            from io import StringIO
            
            structured_data = {
                "accomplishments": [],
                "problems": [],
                "tomorrow_plans": [],
                "metrics": [],
                "departments": set(),
                "dates": [],
                "has_structured_data": False
            }
            
            lines = content.strip().split('\n')
            if not lines or len(lines) < 2:
                print("⚠️ CSV has no data rows")
                return {"has_structured_data": False, "message": "CSV has no data rows"}
            
            # Check if first line contains our headers
            first_line_lower = lines[0].lower()
            has_department = 'department' in first_line_lower
            has_date = 'date' in first_line_lower
            has_problems = 'problem' in first_line_lower or 'challenge' in first_line_lower
            has_achievements = 'accomplish' in first_line_lower or 'achievement' in first_line_lower
            has_plans = 'plan' in first_line_lower or 'future' in first_line_lower or 'tomorrow' in first_line_lower
            
            print(f"Header check: dept={has_department}, date={has_date}, problems={has_problems}, achievements={has_achievements}, plans={has_plans}")
            
            if has_department and (has_problems or has_achievements or has_plans):
                try:
                    csv_file = StringIO(content)
                    csv_reader = csv.DictReader(csv_file)
                    
                    row_count = 0
                    for row in csv_reader:
                        row_count += 1
                        # Process department
                        for key in row:
                            if 'department' in key.lower() and row[key]:
                                dept = row[key].strip()
                                if dept:
                                    structured_data["departments"].add(dept)
                        
                        # Process problems
                        for key in row:
                            if 'problem' in key.lower() or 'challenge' in key.lower():
                                if row[key] and row[key].strip():
                                    problems = self._split_items(row[key])
                                    structured_data["problems"].extend(problems)
                        
                        # Process accomplishments
                        for key in row:
                            if 'accomplish' in key.lower() or 'achievement' in key.lower():
                                if row[key] and row[key].strip():
                                    achievements = self._split_items(row[key])
                                    structured_data["accomplishments"].extend(achievements)
                        
                        # Process future plans
                        for key in row:
                            if 'plan' in key.lower() or 'future' in key.lower() or 'tomorrow' in key.lower():
                                if row[key] and row[key].strip():
                                    plans = self._split_items(row[key])
                                    structured_data["tomorrow_plans"].extend(plans)
                        
                        # Extract metrics from all cells
                        for key, value in row.items():
                            if value and any(word in value.lower() for word in ['%', 'uptime', 'kpi', 'metric', 'score', 'rate']):
                                structured_data["metrics"].append(f"{key}: {value}")
                    
                    print(f"Processed {row_count} rows")
                    
                    # Convert set to list
                    structured_data["departments"] = list(structured_data["departments"])
                    
                    # Check if we found data
                    has_data = any(structured_data[key] for key in ["accomplishments", "problems", "tomorrow_plans", "metrics", "departments"])
                    structured_data["has_structured_data"] = has_data
                    
                    if has_data:
                        # Create formatted content
                        formatted = []
                        formatted.append("=== 5-COLUMN CSV DATA ===")
                        
                        if structured_data["departments"]:
                            formatted.append(f"\nDepartments: {', '.join(structured_data['departments'])}")
                        
                        if structured_data["accomplishments"]:
                            formatted.append("\nACCOMPLISHMENTS:")
                            for acc in structured_data["accomplishments"][:10]:
                                formatted.append(f"- {acc}")
                        
                        if structured_data["problems"]:
                            formatted.append("\nPROBLEMS:")
                            for prob in structured_data["problems"][:10]:
                                formatted.append(f"- {prob}")
                        
                        if structured_data["tomorrow_plans"]:
                            formatted.append("\nFUTURE PLANS:")
                            for plan in structured_data["tomorrow_plans"][:10]:
                                formatted.append(f"- {plan}")
                        
                        if structured_data["metrics"]:
                            formatted.append("\nMETRICS:")
                            for metric in structured_data["metrics"][:5]:
                                formatted.append(f"- {metric}")
                        
                        structured_data["formatted_content"] = "\n".join(formatted)
                    
                    return structured_data
                    
                except Exception as e:
                    print(f"5-column CSV parsing error: {e}")
                    return {"has_structured_data": False, "error": str(e)}
            
            return {"has_structured_data": False, "message": "Doesn't match 5-column format"}
            
        except Exception as e:
            print(f"Error parsing 5-column CSV: {e}")
            return {"has_structured_data": False, "error": str(e)}
    
    def _parse_generic_csv(self, content: str) -> dict:
        """Parse generic CSV files"""
        try:
            from io import StringIO
            
            structured_data = {
                "accomplishments": [],
                "problems": [],
                "tomorrow_plans": [],
                "metrics": [],
                "departments": set(),
                "has_structured_data": False
            }
            
            try:
                csv_file = StringIO(content)
                csv_reader = csv.reader(csv_file)
                rows = list(csv_reader)
                
                if len(rows) > 1:
                    # Extract all text
                    all_text = []
                    for row in rows:
                        for cell in row:
                            if cell and cell.strip():
                                all_text.append(cell.strip())
                    
                    # Create formatted content
                    formatted = ["=== CSV DATA ==="]
                    formatted.extend(all_text[:50])
                    
                    structured_data["formatted_content"] = "\n".join(formatted)
                    
                    # Try to categorize
                    text_content = " ".join(all_text).lower()
                    
                    # Look for departments
                    for row in rows[1:]:  # Skip header
                        for cell in row:
                            if cell and any(dept in cell.lower() for dept in ['it', 'engineering', 'marketing', 'sales', 'hr', 'finance', 'operations', 'support']):
                                structured_data["departments"].add(cell.strip())
                    
                    # Look for accomplishments
                    if any(word in text_content for word in ['completed', 'finished', 'achieved', 'implemented', 'deployed', 'fixed', 'resolved']):
                        # Extract sentences with these words
                        lines = content.split('\n')
                        for line in lines[:20]:
                            line_lower = line.lower()
                            if any(word in line_lower for word in ['completed', 'finished', 'achieved']):
                                structured_data["accomplishments"].append(line[:200])
                    
                    # Look for problems
                    if any(word in text_content for word in ['issue', 'problem', 'error', 'bug', 'failed', 'broken']):
                        lines = content.split('\n')
                        for line in lines[:20]:
                            line_lower = line.lower()
                            if any(word in line_lower for word in ['issue', 'problem', 'error']):
                                structured_data["problems"].append(line[:200])
                    
                    structured_data["departments"] = list(structured_data["departments"])
                    structured_data["has_structured_data"] = True
                
                return structured_data
                
            except Exception as e:
                print(f"Generic CSV parsing error: {e}")
                return {"has_structured_data": False}
            
        except Exception as e:
            print(f"Error parsing generic CSV: {e}")
            return {"has_structured_data": False}
    
    def _process_pdf(self, content_bytes: bytes, result: dict) -> dict:
        """Process PDF files"""
        try:
            pdf_file = io.BytesIO(content_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page_num in range(min(len(pdf_reader.pages), 10)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(page_text)
            
            full_text = "\n".join(text_parts) if text_parts else "PDF (no text extracted)"
            
            result["file_type"] = "pdf"
            result["content"] = full_text[:50000]
            
            # Try to extract structured data
            structured = self._extract_structured_from_text(full_text)
            if structured.get("has_structured_data"):
                result["structured_data"] = structured
            
            return result
            
        except Exception as e:
            result["success"] = False
            result["error"] = str(e)
            return result
    
    def _process_word(self, content_bytes: bytes, result: dict) -> dict:
        """Process Word files"""
        try:
            doc_file = io.BytesIO(content_bytes)
            doc = Document(doc_file)
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            full_text = "\n".join(paragraphs) if paragraphs else "Word document (empty)"
            
            result["file_type"] = "word"
            result["content"] = full_text[:50000]
            
            # Try to extract structured data
            structured = self._extract_structured_from_text(full_text)
            if structured.get("has_structured_data"):
                result["structured_data"] = structured
            
            return result
            
        except Exception as e:
            result["success"] = False
            result["error"] = str(e)
            return result
    
    def _process_excel(self, content_bytes: bytes, result: dict) -> dict:
        """Process Excel files (as text)"""
        try:
            # Try to decode as text first
            try:
                content = content_bytes.decode('utf-8', errors='ignore')
                if len(content) > 100:
                    result["file_type"] = "excel_text"
                    result["content"] = content[:50000]
                    
                    structured = self._extract_structured_from_text(content)
                    if structured.get("has_structured_data"):
                        result["structured_data"] = structured
                    
                    return result
            except:
                pass
            
            # If not text, provide instructions
            result["file_type"] = "excel"
            result["content"] = """📊 Excel File Detected

For best results with Excel files:
1. Save your Excel file as CSV (File → Save As → CSV)
2. Upload the CSV file
3. Or copy-paste the data into the Text Input tab

The AI works best with:
• CSV files with clear headers
• Structured reports in text format
• PDF/Word documents with readable text

You can also try copying your Excel data and using the Text Input option."""
            
            return result
            
        except Exception as e:
            result["success"] = False
            result["error"] = str(e)
            return result
    
    def _process_generic(self, content_bytes: bytes, result: dict, filename_lower: str) -> dict:
        """Process generic/unrecognized files"""
        result["file_type"] = "generic"
        try:
            # Try to decode as text
            content = content_bytes.decode('utf-8', errors='ignore')
            if len(content) > 100:
                result["content"] = content[:50000]
                
                structured = self._extract_structured_from_text(content)
                if structured.get("has_structured_data"):
                    result["structured_data"] = structured
            else:
                result["content"] = f"Binary file: {filename_lower} ({len(content_bytes)} bytes)"
        except:
            result["content"] = f"Binary file: {filename_lower} ({len(content_bytes)} bytes)"
        return result
    
    def _extract_structured_from_text(self, text: str) -> dict:
        """Extract structured data from unstructured text"""
        structured_data = {
            "accomplishments": [],
            "problems": [],
            "tomorrow_plans": [],
            "metrics": [],
            "departments": [],
            "has_structured_data": False
        }
        
        if not text or len(text) < 50:
            return structured_data
        
        text_lower = text.lower()
        lines = text.split('\n')
        
        # Keywords for each category
        accomplishment_keywords = [
            'completed', 'finished', 'achieved', 'implemented', 'deployed',
            'fixed', 'resolved', 'solved', 'delivered', 'launched',
            'created', 'built', 'developed', 'successful', 'done'
        ]
        
        problem_keywords = [
            'issue', 'problem', 'error', 'bug', 'failed', 'broken',
            'slow', 'down', 'outage', 'crash', 'blocked', 'stuck',
            'challenge', 'difficult', 'risk', 'concern'
        ]
        
        plan_keywords = [
            'plan', 'will', 'next', 'tomorrow', 'future', 'scheduled',
            'need to', 'should', 'must', 'going to', 'aim to', 'target',
            'improve', 'enhance', 'optimize', 'upgrade'
        ]
        
        metric_keywords = [
            'uptime', 'kpi', 'metric', 'score', 'rate', 'percentage',
            'efficiency', 'performance', 'response time', 'throughput',
            'accuracy', 'precision', 'recall', 'coverage'
        ]
        
        department_keywords = [
            'it department', 'engineering', 'marketing', 'sales', 
            'human resources', 'hr', 'finance', 'operations', 'support'
        ]
        
        # Extract departments
        departments_found = set()
        for line in lines[:50]:
            line_lower = line.lower()
            for dept in department_keywords:
                if dept in line_lower:
                    # Extract the actual department name
                    for word in line.split():
                        if any(dept_word in word.lower() for dept_word in ['it', 'engineering', 'marketing', 'sales', 'hr', 'finance', 'operations']):
                            departments_found.add(word.strip(',.:;'))
        
        if departments_found:
            structured_data["departments"] = list(departments_found)
        
        # Extract accomplishments
        for line in lines[:100]:
            if len(line.strip()) > 10:
                line_lower = line.lower()
                if any(keyword in line_lower for keyword in accomplishment_keywords):
                    structured_data["accomplishments"].append(line.strip()[:200])
        
        # Extract problems
        for line in lines[:100]:
            if len(line.strip()) > 10:
                line_lower = line.lower()
                if any(keyword in line_lower for keyword in problem_keywords):
                    structured_data["problems"].append(line.strip()[:200])
        
        # Extract plans
        for line in lines[:100]:
            if len(line.strip()) > 10:
                line_lower = line.lower()
                if any(keyword in line_lower for keyword in plan_keywords):
                    structured_data["tomorrow_plans"].append(line.strip()[:200])
        
        # Extract metrics
        for line in lines[:100]:
            if len(line.strip()) > 10:
                line_lower = line.lower()
                if any(keyword in line_lower for keyword in metric_keywords):
                    structured_data["metrics"].append(line.strip()[:200])
        
        # Remove duplicates
        for key in ["accomplishments", "problems", "tomorrow_plans", "metrics", "departments"]:
            if key in structured_data:
                structured_data[key] = list(set(structured_data[key]))[:10]
        
        # Check if we found any data
        has_data = any(structured_data[key] for key in ["accomplishments", "problems", "tomorrow_plans", "metrics", "departments"])
        structured_data["has_structured_data"] = has_data
        
        return structured_data
    
    def _split_items(self, text: str) -> list:
        """Split comma-separated items"""
        if not text:
            return []
        
        # Split by commas, semicolons, or newlines
        items = re.split(r'[,;\n]', text)
        items = [item.strip() for item in items if item.strip()]
        
        return items

# ==================== ENHANCED AI ANALYZER ====================
class EnhancedAIAnalyzer:
    """Enhanced AI analyzer for all file types"""
    
    def __init__(self):
        # Comprehensive word lists
        self.positive_words = {
            'good', 'great', 'excellent', 'outstanding', 'successful',
            'completed', 'finished', 'achieved', 'implemented', 'deployed',
            'fixed', 'resolved', 'solved', 'working', 'improved',
            'positive', 'better', 'fast', 'efficient', 'reliable',
            'stable', 'secure', 'optimized', 'upgraded', 'enhanced',
            'launched', 'delivered', 'created', 'built', 'developed'
        }
        
        self.negative_words = {
            'bad', 'poor', 'failed', 'issue', 'problem', 'error',
            'bug', 'broken', 'slow', 'delayed', 'blocked', 'stuck',
            'negative', 'worse', 'difficult', 'challenge', 'risk',
            'critical', 'urgent', 'emergency', 'outage', 'down',
            'crash', 'defect', 'failure', 'missed', 'incomplete'
        }
        
        self.action_words = {
            'need', 'must', 'should', 'will', 'plan', 'next',
            'tomorrow', 'schedule', 'assign', 'action', 'task',
            'complete', 'finish', 'start', 'implement', 'review',
            'meet', 'discuss', 'prepare', 'update', 'test',
            'deploy', 'fix', 'resolve', 'improve', 'optimize'
        }
        
        self.metric_keywords = {
            'uptime', 'performance', 'efficiency', 'response',
            'throughput', 'accuracy', 'coverage', 'score',
            'rate', 'percentage', 'kpi', 'metric', 'measure'
        }
    
    def analyze(self, text: str, structured_data: dict = None) -> dict:
        """Enhanced AI analysis - FIXED FOR EMPTY TEMPLATES"""
        if not text or len(text.strip()) < 10:
            return self._empty_analysis()
        
        text_lower = text.lower()
        words = re.findall(r'\b\w+\b', text_lower)
        
        # Check if it's just headers or template
        if "this appears to be a template" in text_lower or "empty csv template" in text_lower:
            return self._template_analysis()
        
        # Sentiment analysis
        pos_count = sum(1 for w in words if w in self.positive_words)
        neg_count = sum(1 for w in words if w in self.negative_words)
        
        # Calculate sentiment score
        total_words = len(words)
        sentiment_score = 0
        if total_words > 0:
            sentiment_score = (pos_count - neg_count) / total_words * 100
        
        # Determine sentiment label
        if sentiment_score > 20:
            sentiment_label = "positive"
        elif sentiment_score < -20:
            sentiment_label = "negative"
        else:
            sentiment_label = "neutral"
        
        # Extract topics
        topics = self._extract_topics(text)
        
        # Extract accomplishments, problems, actions, departments
        accomplishments = []
        problems = []
        action_items = []
        departments = []
        
        # Use structured data if available
        if structured_data and structured_data.get("has_structured_data"):
            accomplishments = structured_data.get("accomplishments", [])[:10]
            problems = structured_data.get("problems", [])[:10]
            action_items = structured_data.get("tomorrow_plans", [])[:10]
            metrics = structured_data.get("metrics", [])[:5]
            departments = structured_data.get("departments", [])[:5]
        else:
            # Extract from unstructured text
            lines = text.split('\n')
            
            # Look for section headers
            current_section = None
            section_keywords = {
                'accomplishments': ['accomplishment', 'achievement', 'completed', 'done'],
                'problems': ['problem', 'issue', 'challenge', 'error'],
                'actions': ['plan', 'next', 'action', 'todo', 'task'],
                'metrics': ['metric', 'kpi', 'performance', 'uptime']
            }
            
            for line in lines[:100]:
                line_lower = line.lower().strip()
                line_text = line.strip()
                
                if not line_text:
                    continue
                
                # Check for section headers
                for section, keywords in section_keywords.items():
                    if any(keyword in line_lower for keyword in keywords) and len(line_text) < 100:
                        current_section = section
                        break
                
                # Categorize line based on keywords
                if any(word in line_lower for word in self.positive_words) and len(line_text) > 10:
                    accomplishments.append(line_text[:200])
                
                if any(word in line_lower for word in self.negative_words) and len(line_text) > 10:
                    problems.append(line_text[:200])
                
                if any(word in line_lower for word in self.action_words) and len(line_text) > 10:
                    action_items.append(line_text[:200])
                
                # Extract departments
                if any(dept in line_lower for dept in ['it', 'engineering', 'marketing', 'sales', 'hr', 'finance', 'operations']):
                    # Try to extract department name
                    for word in line_text.split():
                        word_lower = word.lower().strip(',.:;')
                        if word_lower in ['it', 'engineering', 'marketing', 'sales', 'hr', 'finance', 'operations']:
                            departments.append(word.capitalize())
        
        # Remove duplicates
        accomplishments = list(set(accomplishments))[:10]
        problems = list(set(problems))[:10]
        action_items = list(set(action_items))[:10]
        departments = list(set(departments))[:5]
        
        # Extract metrics from text
        metrics = self._extract_metrics(text)
        
        # Generate summary
        summary = self._generate_summary(text)
        
        # Generate conclusion
        conclusion = self._generate_conclusion(sentiment_label, sentiment_score, accomplishments, problems, action_items, metrics, departments)
        
        return {
            "sentiment": {
                "label": sentiment_label,
                "score": round(sentiment_score, 2),
                "positive_words": pos_count,
                "negative_words": neg_count
            },
            "topics": topics[:5],
            "summary": summary,
            "word_count": len(words),
            "accomplishments": accomplishments,
            "problems": problems,
            "action_items": action_items,
            "metrics": metrics,
            "departments": departments,
            "ai_conclusion": conclusion,
            "analysis_complete": True,
            "timestamp": datetime.now().isoformat()
        }
    
    def _template_analysis(self) -> dict:
        """Special analysis for empty templates"""
        return {
            "sentiment": {"label": "neutral", "score": 0, "positive_words": 0, "negative_words": 0},
            "topics": ["template", "csv", "format"],
            "summary": "CSV template detected. Please add your data to the rows below the headers.",
            "word_count": 0,
            "accomplishments": [],
            "problems": ["No data in CSV - this appears to be a template"],
            "action_items": ["1. Fill the template with your data", "2. Save as CSV", "3. Upload again"],
            "metrics": [],
            "departments": [],
            "ai_conclusion": {
                "text": """📋 **TEMPLATE DETECTED**

This appears to be an empty CSV template.

✅ **NEXT STEPS:**
1. Open this file in Excel or Google Sheets
2. Fill in your data under the headers:
   - Department: e.g., "IT", "Marketing"
   - Date: e.g., "2024-01-15"
   - Problems: e.g., "Network issue, VPN down"
   - Achievements: e.g., "Fixed login bug"
   - Future_Plan: e.g., "Schedule maintenance"

3. Save the file as CSV
4. Upload the filled file for analysis

💡 **TIP:** You can copy the sample data from the "Sample CSV" button to get started!""",
                "timestamp": datetime.now().isoformat()
            },
            "analysis_complete": False,
            "is_template": True
        }
    
    def _extract_topics(self, text: str) -> list:
        """Extract main topics from text"""
        words = re.findall(r'\b[a-z]{4,}\b', text.lower())
        
        # Common words to exclude
        stop_words = {
            'that', 'with', 'this', 'from', 'have', 'they', 'what',
            'were', 'when', 'your', 'about', 'would', 'there',
            'their', 'which', 'could', 'other', 'these', 'some',
            'more', 'most', 'also', 'into', 'only', 'over', 'after',
            'even', 'much', 'such', 'than', 'them', 'then', 'they',
            'well', 'were', 'what', 'when', 'whom', 'will', 'with',
            'your', 'said', 'like', 'just'
        }
        
        word_counts = Counter()
        for word in words:
            if word not in stop_words:
                word_counts[word] += 1
        
        # Get top topics
        topics = [word for word, count in word_counts.most_common(10)]
        
        return topics
    
    def _extract_metrics(self, text: str) -> list:
        """Extract metrics/KPIs from text"""
        metrics = []
        lines = text.split('\n')
        
        metric_patterns = [
            r'(\d+\.?\d*)\s*%',  # percentages
            r'(\d+\.?\d*)\s*(ms|s|sec|seconds)',  # time
            r'(\d+\.?\d*)\s*(mb|gb|tb)',  # size
            r'uptime.*?(\d+\.?\d*)\s*%',  # uptime
            r'response.*?(\d+\.?\d*)',  # response time
            r'accuracy.*?(\d+\.?\d*)\s*%',  # accuracy
        ]
        
        for line in lines[:50]:
            line_lower = line.lower()
            
            # Check for metric keywords
            if any(keyword in line_lower for keyword in self.metric_keywords):
                metrics.append(line.strip()[:150])
            
            # Check for patterns
            for pattern in metric_patterns:
                matches = re.search(pattern, line_lower)
                if matches:
                    metrics.append(line.strip()[:150])
                    break
        
        return list(set(metrics))[:5]
    
    def _generate_summary(self, text: str) -> str:
        """Generate a concise summary"""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        if not lines:
            return "No content to summarize"
        
        # Use first meaningful line as summary
        for line in lines:
            if len(line) > 20 and len(line) < 200:
                return line + "..."
        
        # Fallback
        return lines[0][:150] + "..." if len(lines[0]) > 150 else lines[0]
    
    def _generate_conclusion(self, sentiment, score, accomplishments, problems, actions, metrics, departments):
        """Generate AI conclusion - FIXED TO INCLUDE DEPARTMENTS"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        conclusion_parts = []
        
        # Overall assessment
        conclusion_parts.append("🤖 **AI ANALYSIS REPORT**")
        conclusion_parts.append(f"**Generated**: {timestamp}")
        conclusion_parts.append("")
        
        if departments:
            conclusion_parts.append(f"**Departments**: {', '.join(departments)}")
            conclusion_parts.append("")
        
        if sentiment == "positive" and score > 40:
            conclusion_parts.append("✅ **OVERALL STATUS: EXCELLENT**")
            conclusion_parts.append("- Strong positive performance reported")
            conclusion_parts.append("- Team is exceeding expectations")
        elif sentiment == "positive":
            conclusion_parts.append("✅ **OVERALL STATUS: GOOD**")
            conclusion_parts.append("- Positive progress observed")
            conclusion_parts.append("- Team is meeting objectives")
        elif sentiment == "negative" or score < -20:
            conclusion_parts.append("⚠️ **OVERALL STATUS: NEEDS ATTENTION**")
            conclusion_parts.append("- Issues identified require focus")
            conclusion_parts.append("- Immediate review recommended")
        else:
            conclusion_parts.append("ℹ️ **OVERALL STATUS: STABLE**")
            conclusion_parts.append("- Standard operations reported")
            conclusion_parts.append("- Continue monitoring progress")
        
        conclusion_parts.append("")
        
        # Key findings
        if accomplishments:
            conclusion_parts.append("🏆 **KEY ACHIEVEMENTS:**")
            for i, acc in enumerate(accomplishments[:5], 1):
                conclusion_parts.append(f"{i}. {acc}")
            conclusion_parts.append("")
        
        if problems:
            conclusion_parts.append("🔴 **KEY CHALLENGES:**")
            for i, prob in enumerate(problems[:5], 1):
                conclusion_parts.append(f"{i}. {prob}")
            conclusion_parts.append("")
        
        # Recommendations
        conclusion_parts.append("💡 **RECOMMENDATIONS:**")
        
        if problems:
            conclusion_parts.append("1. **Prioritize Issue Resolution**")
            conclusion_parts.append(f"   - Address the {len(problems)} identified issues")
        
        if actions:
            conclusion_parts.append("2. **Execute Action Plan**")
            conclusion_parts.append(f"   - Follow through on {len(actions)} planned actions")
        else:
            conclusion_parts.append("2. **Develop Action Plan**")
            conclusion_parts.append("   - Create specific tasks for next period")
        
        if metrics:
            conclusion_parts.append("3. **Monitor Performance Metrics**")
            conclusion_parts.append("   - Track key metrics for continuous improvement")
        
        conclusion_parts.append("")
        conclusion_parts.append("📊 **ANALYSIS CONFIDENCE**: High")
        conclusion_parts.append("🚀 **AI ENGINE**: DeepSeek v2.2")
        
        return {
            "text": "\n".join(conclusion_parts),
            "timestamp": timestamp
        }
    
    def _empty_analysis(self) -> dict:
        return {
            "sentiment": {"label": "neutral", "score": 0, "positive_words": 0, "negative_words": 0},
            "topics": [],
            "summary": "No meaningful content detected",
            "word_count": 0,
            "accomplishments": [],
            "problems": [],
            "action_items": [],
            "metrics": [],
            "departments": [],
            "ai_conclusion": {
                "text": "⚠️ Insufficient content for meaningful analysis.\n\nPlease provide a detailed report with accomplishments, challenges, and plans for comprehensive insights.",
                "timestamp": datetime.now().isoformat()
            },
            "analysis_complete": False
        }

# ==================== INITIALIZE ====================
file_processor = EnhancedFileProcessor()
ai_analyzer = EnhancedAIAnalyzer()

# conn = sqlite3.connect('reports.db', check_same_thread=False)
#     cursor = conn.cursor()
    
#     reports = execute_query("SELECT * FROM reports", fetch_all=True)
    
#     # For complex queries with transactions:

# ==================== DATABASE ====================
# 
def setup_db():
    """Create database tables if they don't exist"""
    conn = None
    cursor = None
    
    try:
        conn = get_connection()
        cursor = conn.cursor()
        
        if USE_SQLITE:
            # SQLite: Check if table exists
            cursor.execute("""
                SELECT name FROM sqlite_master 
                WHERE type='table' AND name='reports'
            """)
            
            if not cursor.fetchone():
                print("📦 Creating SQLite table...")
                cursor.execute("""
                    CREATE TABLE reports (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        department TEXT,
                        report_date TEXT,
                        filename TEXT,
                        file_type TEXT,
                        file_size INTEGER,
                        content_preview TEXT,
                        word_count INTEGER,
                        upload_date TEXT,
                        ai_analysis TEXT,
                        ai_conclusion TEXT,
                        total_reports INTEGER DEFAULT 0,
                        total_problems INTEGER DEFAULT 0,
                        total_achievements INTEGER DEFAULT 0,
                        sentiment_label TEXT DEFAULT 'neutral'
                    )
                """)
                
                # Create indexes
                try:
                    cursor.execute("CREATE INDEX idx_reports_department ON reports(department)")
                    cursor.execute("CREATE INDEX idx_reports_upload_date ON reports(upload_date)")
                    cursor.execute("CREATE INDEX idx_reports_sentiment ON reports(sentiment_label)")
                except:
                    pass  # Indexes might fail if they exist
                
                print("✅ SQLite database created")
            else:
                print("✅ SQLite database already exists")
        else:
            # PostgreSQL: Check if table exists
            cursor.execute("""
                SELECT EXISTS (
                    SELECT FROM information_schema.tables 
                    WHERE table_name = 'reports'
                )
            """)
            
            if not cursor.fetchone()["exists"]:
                print("📦 Creating PostgreSQL table...")
                cursor.execute("""
                    CREATE TABLE reports (
                        id SERIAL PRIMARY KEY,
                        department TEXT,
                        report_date TEXT,
                        filename TEXT,
                        file_type TEXT,
                        file_size INTEGER,
                        content_preview TEXT,
                        word_count INTEGER,
                        upload_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        ai_analysis JSONB,
                        ai_conclusion JSONB,
                        total_reports INTEGER DEFAULT 0,
                        total_problems INTEGER DEFAULT 0,
                        total_achievements INTEGER DEFAULT 0,
                        sentiment_label TEXT DEFAULT 'neutral'
                    )
                """)
                
                # Create indexes
                cursor.execute("CREATE INDEX idx_reports_department ON reports(department)")
                cursor.execute("CREATE INDEX idx_reports_upload_date ON reports(upload_date)")
                cursor.execute("CREATE INDEX idx_reports_sentiment ON reports(sentiment_label)")
                
                print("✅ PostgreSQL database created")
            else:
                print("✅ PostgreSQL database already exists")
        
        conn.commit()
        
    except Exception as e:
        print(f"❌ Database setup error: {e}")
        if conn:
            conn.rollback()
        raise
        
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

# ==================== API ENDPOINTS ====================
@app.get("/")
def home():
    return {
        "message": "✅ Fixed Report Analyzer AI", 
        "status": "running",
        "version": "2.2",
        "features": [
            "PDF, Word, Excel, CSV, Text support",
            "Enhanced AI analysis",
            "Performance dashboard",
            "Template download fixed",
            "Department extraction"
        ]
    }

@app.get("/api/")
def api_root():
    return {"message": "Report Analyzer API", "status": "active", "version": "2.2"}

# FIXED: Added this endpoint for frontend
@app.get("/api/template")
def get_template():
    """Alias for template download - FIXED FOR FRONTEND"""
    return download_template("csv")

@app.post("/api/upload-csv")
async def upload_csv_endpoint(
    department: str = Form(...), 
    date: str = Form(...), 
    file: UploadFile = File(...)
):
    """Alias for /api/upload for frontend compatibility"""
    return await upload_report(department, date, file)

@app.get("/api/test-data")
def test_data():
    """Provide test data for frontend - UPDATED WITH DEPARTMENTS"""
    return {
        "success": True,
        "summary": {
            "total_reports": 3,
            "total_problems": 6,
            "total_achievements": 5,
            "total_future_plans": 4,
            "departments": ["IT", "Engineering", "Support"],
            "common_problems": [["Database slow", 2], ["VPN issue", 1], ["Network down", 1]],
            "common_achievements": [["Fixed login bug", 2], ["Deployed dashboard", 1]],
            "common_plans": [["Schedule maintenance", 2], ["Optimize queries", 1]]
        },
        "ai_analysis": {
            "insights": [
                "📊 Departments: IT, Engineering, Support",
                "🔴 Top Problems: Database slow (2), VPN issue (1)",
                "✅ Top Achievements: Fixed login bug (2), Deployed dashboard (1)"
            ],
            "recommendations": [
                "🚨 Priority: Address 'Database slow' - mentioned 2 times",
                "📋 Action: Schedule a team meeting to address recurring issues",
                "📈 Monitor: Track progress on identified issues weekly"
            ],
            "sentiment": "Positive",
            "summary": {
                "total_reports": 3,
                "total_problems": 6,
                "total_achievements": 5,
                "total_plans": 4
            },
            "departments": ["IT", "Engineering", "Support"]
        }
    }

@app.post("/api/upload")
async def upload_report(
    department: str = Form(...), 
    date: str = Form(...), 
    file: UploadFile = File(...)
):
    try:
        print(f"📤 Uploading: {file.filename} for department: {department}")
        
        content_bytes = await file.read()
        
        # Process file
        processing_result = file_processor.process_file(content_bytes, file.filename)
        
        if not processing_result["success"]:
            return JSONResponse(
                status_code=400,
                content={"success": False, "error": processing_result["error"]}
            )
        
        print(f"📄 File processed as: {processing_result['file_type']}")
        
        # AI analysis
        structured_data = processing_result.get("structured_data", {})
        ai_result = ai_analyzer.analyze(processing_result["content"], structured_data)
        
        # Extract departments from CSV data if available
        csv_departments = []
        if structured_data and structured_data.get("departments"):
            csv_departments = structured_data.get("departments", [])
        
        # Combine departments: from form input and from CSV data
        all_departments = list(set([department] + csv_departments))
        
        # Calculate stats
        total_reports = 1
        total_problems = len(ai_result.get("problems", []))
        total_achievements = len(ai_result.get("accomplishments", []))
        
        # Save to database
        conn = sqlite3.connect('reports.db', check_same_thread=False)
        cursor = conn.cursor()
        
        cursor.execute("""
    INSERT INTO reports 
    (department, report_date, filename, file_type, file_size, 
     content_preview, word_count, upload_date, ai_analysis, ai_conclusion,
     total_reports, total_problems, total_achievements, sentiment_label)
    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
""", (
    department, date, file.filename, processing_result["file_type"],
    processing_result["file_size"], processing_result["content"][:500],
    ai_result["word_count"], datetime.now().isoformat(),
    json.dumps(ai_result), json.dumps(ai_result["ai_conclusion"]),
    1,  # total_reports
    len(ai_result.get("problems", [])),  # total_problems
    len(ai_result.get("accomplishments", [])),  # total_achievements
    ai_result["sentiment"]["label"]  # sentiment_label
))
        
        conn.commit()
        conn.close()
        
        # Prepare summary with departments
        summary_data = {
            "total_reports": 1,
            "total_problems": len(ai_result.get("problems", [])),
            "total_achievements": len(ai_result.get("accomplishments", [])),
            "total_future_plans": len(ai_result.get("action_items", [])),
            "departments": all_departments
        }
        
        # Add common problems and achievements if available
        if ai_result.get("problems"):
            summary_data["common_problems"] = [[prob, 1] for prob in ai_result["problems"][:5]]
        
        if ai_result.get("accomplishments"):
            summary_data["common_achievements"] = [[acc, 1] for acc in ai_result["accomplishments"][:5]]
        
        if ai_result.get("action_items"):
            summary_data["common_plans"] = [[plan, 1] for plan in ai_result["action_items"][:5]]
        
        return {
            "success": True,
            "message": f"✅ {processing_result['file_type'].upper()} file analyzed successfully!",
            "department": department,
            "file_info": {
                "filename": file.filename,
                "type": processing_result["file_type"],
                "size": processing_result["file_size"],
                "structured": bool(structured_data.get("has_structured_data", False))
            },
            "ai_analysis": ai_result,
            "ai_conclusion": ai_result["ai_conclusion"],
            "summary": summary_data
        }
        
    except Exception as e:
        print(f"❌ Upload error: {traceback.format_exc()}")
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )

@app.post("/api/analyze-text")
async def analyze_text(text: str = Form(...)):
    """Analyze text directly"""
    try:
        if not text or len(text.strip()) < 10:
            return {"success": False, "error": "Text too short (minimum 10 characters)"}
        
        ai_result = ai_analyzer.analyze(text)
        
        # Prepare summary with departments
        summary_data = {
            "total_reports": 1,
            "total_problems": len(ai_result.get("problems", [])),
            "total_achievements": len(ai_result.get("accomplishments", [])),
            "total_future_plans": len(ai_result.get("action_items", [])),
            "departments": ai_result.get("departments", [])
        }
        
        # Add common problems and achievements if available
        if ai_result.get("problems"):
            summary_data["common_problems"] = [[prob, 1] for prob in ai_result["problems"][:5]]
        
        if ai_result.get("accomplishments"):
            summary_data["common_achievements"] = [[acc, 1] for acc in ai_result["accomplishments"][:5]]
        
        if ai_result.get("action_items"):
            summary_data["common_plans"] = [[plan, 1] for plan in ai_result["action_items"][:5]]
        
        return {
            "success": True,
            "message": "✅ Text analyzed with AI!",
            "word_count": len(text.split()),
            "ai_analysis": ai_result,
            "ai_conclusion": ai_result["ai_conclusion"],
            "summary": summary_data
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.get("/api/reports")
def get_reports(limit: int = 20, offset: int = 0):
    try:
        conn = sqlite3.connect('reports.db', check_same_thread=False)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        cursor.execute('SELECT * FROM reports ORDER BY id DESC LIMIT ? OFFSET ?', (limit, offset))
        
        reports = []
        for row in cursor.fetchall():
            report = dict(row)
            
            # Parse JSON fields
            for field in ['ai_analysis', 'ai_conclusion']:
                if report.get(field):
                    try:
                        report[field] = json.loads(report[field])
                    except:
                        report[field] = None
            
            reports.append(report)
        
        # Get total count
        cursor.execute('SELECT COUNT(*) FROM reports')
        total = cursor.fetchone()[0]
        
        conn.close()
        
        return {
            "reports": reports, 
            "count": len(reports),
            "total": total,
            "limit": limit,
            "offset": offset
        }
        
    except Exception as e:
        return {"error": str(e), "reports": []}

@app.get("/api/stats")
def get_stats():
    try:
        conn = sqlite3.connect('reports.db', check_same_thread=False)
        cursor = conn.cursor()
        
        # Basic stats
        cursor.execute('SELECT COUNT(*) FROM reports')
        total_reports = cursor.fetchone()[0]
        
        cursor.execute('SELECT COUNT(DISTINCT department) FROM reports')
        total_departments = cursor.fetchone()[0]
        
        # Today's reports
        today = datetime.now().strftime("%Y-%m-%d")
        cursor.execute('SELECT COUNT(*) FROM reports WHERE DATE(upload_date) = DATE(?)', (today,))
        today_reports = cursor.fetchone()[0]
        
        # Sentiment distribution
        cursor.execute('SELECT sentiment_label, COUNT(*) FROM reports GROUP BY sentiment_label')
        sentiment_distribution = dict(cursor.fetchall())
        
        # Aggregate problems and achievements
        cursor.execute('SELECT SUM(total_problems), SUM(total_achievements) FROM reports')
        sum_result = cursor.fetchone()
        total_problems = sum_result[0] or 0
        total_achievements = sum_result[1] or 0
        
        conn.close()
        
        return {
            "total_reports": total_reports,
            "total_departments": total_departments,
            "today_reports": today_reports,
            "sentiment_distribution": sentiment_distribution,
            "total_problems": total_problems,
            "total_achievements": total_achievements,
            "ai_enabled": True,
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        return {"error": str(e)}

@app.get("/api/performance-dashboard")
def performance_dashboard(days: int = 30):
    """Performance dashboard data"""
    try:
        since_date = (datetime.now() - timedelta(days=days)).strftime("%Y-%m-%d")
        
        conn = sqlite3.connect('reports.db', check_same_thread=False)
        cursor = conn.cursor()
        
        # Get reports by date
        cursor.execute('''
            SELECT DATE(upload_date) as date, 
                   COUNT(*) as count,
                   SUM(CASE WHEN sentiment_label = 'positive' THEN 1 ELSE 0 END) as positive,
                   SUM(CASE WHEN sentiment_label = 'negative' THEN 1 ELSE 0 END) as negative
            FROM reports 
            WHERE DATE(upload_date) >= DATE(?)
            GROUP BY DATE(upload_date)
            ORDER BY date DESC
        ''', (since_date,))
        
        trends = {}
        for row in cursor.fetchall():
            date, count, positive, negative = row
            trends[date] = {
                "total": count,
                "positive": positive or 0,
                "negative": negative or 0
            }
        
        # Department performance
        cursor.execute('''
            SELECT department, 
                   COUNT(*) as report_count,
                   AVG(CASE WHEN sentiment_label = 'positive' THEN 1 
                            WHEN sentiment_label = 'negative' THEN -1 
                            ELSE 0 END) as sentiment_score
            FROM reports 
            WHERE DATE(upload_date) >= DATE(?)
            GROUP BY department
        ''', (since_date,))
        
        departments = {}
        for row in cursor.fetchall():
            dept, count, score = row
            departments[dept] = {
                "report_count": count,
                "sentiment_score": score or 0
            }
        
        conn.close()
        
        # Generate recommendations
        recommendations = []
        total_recent_reports = sum(t.get("total", 0) for t in trends.values())
        
        if total_recent_reports == 0:
            recommendations.append("N o recent reports found. Start by uploading some reports.")
        else:
            # Calculate overall sentiment
            total_positive = sum(t["positive"] for t in trends.values())
            total_negative = sum(t["negative"] for t in trends.values())
            
            if total_negative > total_positive:
                recommendations.append("⚠️ **Focus on Issue Resolution**: Negative reports exceed positive ones. Prioritize addressing identified problems.")
            
            if len(departments) < 3:
                recommendations.append("📊 **Expand Department Coverage**: Consider involving more departments for comprehensive insights.")
            
            recommendations.append("📈 **Regular Monitoring**: Continue tracking performance metrics weekly.")
            recommendations.append("🤝 **Cross-Department Collaboration**: Share best practices between teams.")
        
        return {
            "success": True,
            "total_reports": sum(t.get("total", 0) for t in trends.values()),
            "trends": trends,
            "departments": departments,
            "recommendations": recommendations,
            "period_days": days
        }
        
    except Exception as e:
        print(f"Dashboard error: {e}")
        return {"success": False, "error": str(e)}

@app.get("/api/download-template/{template_type}")
def download_template(template_type: str = "csv"):
    """Download report templates - FIXED FOR FRONTEND"""
    if template_type.lower() == "excel":
        template = """# EXCEL TEMPLATE FOR REPORT ANALYZER AI

**Save this file as .xlsx or .csv**

Recommended structure:
1. Department
2. Date
3. Problems/Challenges
4. Accomplishments/Achievements  
5. Plans for Tomorrow/Future Actions

Example format:
| Department | Date       | Problems                     | Achievements                  | Plans for Tomorrow           |
|------------|------------|------------------------------|-------------------------------|------------------------------|
| IT         | 2024-01-15 | Network issue, VPN down      | Fixed login bug               | Optimize queries             |
| IT         | 2024-01-15 | Database slow                | Deployed dashboard            | Schedule maintenance         |

Instructions:
1. Open Excel
2. Create columns as shown above
3. Enter your data
4. Save as CSV file
5. Upload for AI analysis

Alternative: Use the CSV template for best results."""
        
        return PlainTextResponse(
            content=template,
            headers={"Content-Disposition": "attachment; filename=excel_template.txt"},
            media_type="text/plain"
        )
    
    elif template_type.lower() == "detailed":
        template = """# DETAILED REPORT TEMPLATE
# Use this for comprehensive reports

## Department: [Your Department]

## Report Date: [YYYY-MM-DD]

## 1. ACCOMPLISHMENTS (What went well?)
- 
- 
- 

## 2. PROBLEMS/CHALLENGES (What issues did you encounter?)
- 
- 
- 

## 3. METRICS/KPIs (Performance indicators)
- Uptime: 
- Response time: 
- Other metrics: 

## 4. PLANS FOR TOMORROW/NEXT PERIOD
- 
- 
- 

## 5. ADDITIONAL NOTES
- 
- 

**Instructions:**
1. Fill in each section
2. Be specific and detailed
3. Save as .txt or .docx file
4. Upload for AI analysis

**Tips:**
- Use bullet points for clarity
- Include specific numbers where possible
- Mention team members if relevant
- Note dependencies or blockers"""
        
        return PlainTextResponse(
            content=template,
            headers={"Content-Disposition": "attachment; filename=detailed_template.txt"},
            media_type="text/plain"
        )
    
    else:
        # Default CSV template - FIXED: Now includes sample data
        template = """Department,Date,Problems,Achievements,Future_Plan
IT,2024-01-15,"Network issue, VPN down","Fixed login bug, Deployed dashboard","Optimize queries, Schedule maintenance"
IT,2024-01-15,"Database slow, Backup outdated","Root cause identified, Monitoring added","Upgrade backup system, Write documentation"
Engineering,2024-01-15,"Bug in API v2, Test coverage low","Project phase completed, Team trained","Client meeting, Fix API bug"
Support,2024-01-15,"High ticket volume, Long response time","Implemented chatbot, Created FAQ","Hire more staff, Automate responses"

✅ **5-COLUMN CSV FORMAT:**
• Each ROW = one report entry
• Each COLUMN = category
• Multiple items: separate with commas

📊 **AI WILL ANALYZE:**
• Departments from all rows
• Accomplishments from all rows
• Problems from all rows  
• Plans from all rows
• Overall trends and patterns

**How to use:**
1. Copy this template
2. Replace with your data
3. Save as .csv file
4. Upload for analysis"""
        
        return PlainTextResponse(
            content=template,
            headers={"Content-Disposition": "attachment; filename=report_template.csv"},
            media_type="text/csv"
        )

@app.get("/api/health")
def health_check():
    return {
        "status": "healthy", 
        "timestamp": datetime.now().isoformat(),
        "ai": "DeepSeek AI Active",
        "version": "2.2",
        "features": ["PDF", "Word", "Excel", "CSV", "Text"]
    }

# ==================== FRONTEND SERVING ====================

# Add these to your existing app.py, in the FRONTEND SERVING section

@app.get("/ui/dashboard-analytics")
async def serve_dashboard_analytics():
    try:
        with open("frontend/dashboard.html", "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        return HTMLResponse("Analytics dashboard page not found")

@app.get("/ui/dashboard")
async def serve_dashboard():
    try:
        with open("frontend/index.html", "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        return HTMLResponse("Dashboard page not found")

@app.get("/ui/upload")
async def serve_upload():
    try:
        with open("frontend/upload.html", "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        return HTMLResponse("Upload page not found")

@app.get("/ui/reports")
async def serve_reports():
    try:
        with open("frontend/reports.html", "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        return HTMLResponse("Reports page not found")

@app.get("/ui/analysis")
async def serve_analysis():
    try:
        with open("frontend/analysis.html", "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        return HTMLResponse("Analysis page not found")

@app.get("/ui", response_class=HTMLResponse)
async def serve_frontend():
    # Read the HTML file you provided
    try:
        with open("frontend/index.html", "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        # Return a simple page if file not found
        return HTMLResponse("""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Report Analyzer AI</title>
            <style>
                body { font-family: Arial, sans-serif; padding: 40px; text-align: center; }
                h1 { color: #333; }
                .status { background: #10b981; color: white; padding: 10px; border-radius: 5px; display: inline-block; }
            </style>
        </head>
        <body>
            <h1>✅ Enhanced Report Analyzer AI - Backend Running</h1>
            <p class="status">Supports: PDF, Word, Excel, CSV, Text</p>
            <p><a href="/docs">API Documentation</a></p>
        </body>
        </html>
        """)

# ==================== RUN APP ====================
if __name__ == "__main__":
    import uvicorn
    
    print("=" * 60)
    print("🚀 FIXED REPORT ANALYZER AI v2.2")
    print("=" * 60)
    print("📁 Supported Formats: PDF, Word, Excel, CSV, Text")
    print("🤖 AI: DeepSeek Enhanced Analysis")
    print("✅ FIXED: Template download and department extraction")
    print("📡 API: http://localhost:8000")
    print("📚 Docs: http://localhost:8000/docs")
    print("🌐 UI: http://localhost:8000/ui")
    print("=" * 60)
    
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)